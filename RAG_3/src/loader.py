import pymupdf
import docx
import os

def load_pdf(pdf_path):
    """
    Load theo tung trang cua pdf
    """
    doc = pymupdf.open(pdf_path)
    pages = []
    for page in doc:
        pages.append(
            page.get_text()
        )
    return pages

def load_docx(docx_path):
    """
    Load van ban tu file Word (.docx)
    """
    doc = docx.Document(docx_path)
    text_blocks = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_blocks.append(para.text.strip())
    # Doc them du lieu tu bang (tables) neu co
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_blocks.append(cell.text.strip())
    return text_blocks

def load_document(file_path):
    """
    Tu dong nhan dang dinh dang file de doc (.pdf hoac .docx)
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    else:
        raise ValueError(f"Dinh dang file khong duoc ho tro: {ext}")